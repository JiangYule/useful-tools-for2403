import docx
from tkinter import messagebox
from time import localtime
import os
import keyboard

file = docx.Document("晚自习作业.docx")
file.add_page_break()
s = localtime()
file.add_paragraph(str(s.tm_mon) + '月' + str(s.tm_mday) + '日作业')
file.add_paragraph('语文:')
file.add_paragraph('数学:')
file.add_paragraph('英语:')
file.add_paragraph('道法:')
file.add_paragraph('历史:')
file.add_paragraph('地理:')
keyboard.press_and_release("f11")
keyboard.press_and_release('ctrl+end')
file.add_paragraph('test')
try:
    file.save("晚自习作业.docx")
    messagebox.showinfo(title="成功", message="新建完成")
    os.system("start 晚自习作业.docx")
except PermissionError:
    messagebox.showerror(title="错误", message="请关闭其他打开作业文件的软件再重试!")

