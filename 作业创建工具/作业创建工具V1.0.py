from docx import Document
from tkinter import messagebox
from time import localtime
import os
import json


def default_settings():
    config = {"subjects": ["语文", "数学", "英语", "道法", "历史", "地理"], "path": "example.docx"}
    with open("settings.json", "w", encoding="utf-8") as fw:
        json.dump(config, fw, ensure_ascii=False, indent=4)


# default_settings()
with open("settings.json", "r", encoding="utf-8") as fr:
    data = json.load(fr)
file = Document(data["path"])
file.add_page_break()
s = localtime()
file.add_paragraph(str(s.tm_mon) + '月' + str(s.tm_mday) + '日作业')
for sub in data["subjects"]:
    file.add_paragraph(sub + ":")
try:
    file.save(data["path"])
    messagebox.showinfo(title="成功", message="新建完成")
    os.system("start" + data["path"])
except PermissionError:
    messagebox.showerror(title="错误", message="请关闭其他打开作业文件的软件再重试!")
